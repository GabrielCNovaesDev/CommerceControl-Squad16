"""
Gera ENTREGA_FINAL_SQUAD_16.docx a partir do Markdown,
usando o template original (.docx) como base para preservar
fontes (Montserrat), header/footer com imagens e estilos.
"""
import re
import shutil
import zipfile
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ── Paths ────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent  # CommerceControl-Squad16/
TEMPLATE_DIR = Path(__file__).resolve().parent / "docx_template" / "original"
TEMPLATE_DOCX = TEMPLATE_DIR / "ENTREGA_FINAL_SQUAD_16.docx"  # vamos copiar
SOURCE_DOCX = ROOT.parent / "DOCUMENTAÇÃO" / "ENTREGA_FINAL_SQUAD_16.docx"
OUTPUT_DIR = ROOT.parent / "DOCUMENTAÇÃO"
OUTPUT_PATH = OUTPUT_DIR / "ENTREGA_FINAL_SQUAD_16.docx"
MD_PATH = ROOT / "ENTREGA_FINAL_SQUAD_16.md"


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_code_block(doc, code_text):
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F4F4F4')
        p_pr.append(shd)


def parse_inline(text):
    segments = []
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            segments.append((text[pos:match.start()], False, False, False))
        token = match.group(0)
        if token.startswith('**'):
            segments.append((token[2:-2], True, False, False))
        elif token.startswith('*'):
            segments.append((token[1:-1], False, True, False))
        elif token.startswith('`'):
            segments.append((token[1:-1], False, False, True))
        pos = match.end()
    if pos < len(text):
        segments.append((text[pos:], False, False, False))
    return segments


def add_formatted_runs(paragraph, text, base_size=11):
    for seg_text, is_bold, is_italic, is_code in parse_inline(text):
        if not seg_text:
            continue
        run = paragraph.add_run(seg_text)
        run.font.name = 'Arial'
        run.font.size = Pt(base_size)
        if is_bold:
            run.bold = True
        if is_italic:
            run.italic = True
        if is_code:
            run.font.name = 'Consolas'
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)


def add_table_from_md(doc, lines, start_idx):
    header_line = lines[start_idx]
    headers = [c.strip() for c in header_line.strip('|').split('|')]
    rows = []
    i = start_idx + 2
    while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip() != '':
        row_cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
        rows.append(row_cells)
        i += 1

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True

    # Adiciona bordas a todas as células manualmente
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

    hdr = table.rows[0]
    for j, h in enumerate(headers):
        cell = hdr.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        add_formatted_runs(p, h, base_size=10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "434343")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Cabeçalho em branco (será sobrescrito abaixo)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for j, cell_text in enumerate(row):
            if j >= len(headers):
                break
            cell = tr.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            add_formatted_runs(p, cell_text, base_size=10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F8F8F8")

    return i


def process_markdown(md_text, doc):
    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('```'):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                in_code = False
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        m = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt({1: 18, 2: 14, 3: 10, 4: 8, 5: 6, 6: 4}.get(level, 6))
            p.paragraph_format.space_after = Pt({1: 10, 2: 8, 3: 6, 4: 4, 5: 4, 6: 2}.get(level, 4))
            run = p.add_run(text)
            if level <= 2:
                run.font.name = 'Montserrat'
            else:
                run.font.name = 'Arial'
            run.font.size = Pt({1: 22, 2: 16, 3: 13, 4: 11, 5: 11, 6: 11}.get(level, 11))
            if level == 1:
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                run.bold = True
            elif level == 2:
                run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
                run.bold = True
            else:
                run.font.color.rgb = RGBColor(0x43, 0x43, 0x43) if level <= 3 else RGBColor(0x66, 0x66, 0x66)
                if level <= 3:
                    run.bold = True
            i += 1
            continue

        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        if stripped.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|', lines[i + 1].strip()):
            i = add_table_from_md(doc, lines, i)
            continue

        m = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_after = Pt(4)
            num_run = p.add_run(f"{m.group(1)}. ")
            num_run.font.name = 'Arial'
            num_run.font.size = Pt(11)
            num_run.bold = True
            add_formatted_runs(p, m.group(2), base_size=11)
            i += 1
            continue

        m = re.match(r'^[-*]\s+(.+)$', stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_after = Pt(4)
            bullet_run = p.add_run("• ")
            bullet_run.font.name = 'Arial'
            bullet_run.font.size = Pt(11)
            bullet_run.bold = True
            add_formatted_runs(p, m.group(1), base_size=11)
            i += 1
            continue

        m = re.match(r'^\s*-\s+\[ \]\s+(.+)$', stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_after = Pt(4)
            check_run = p.add_run("☐  ")
            check_run.font.name = 'Arial'
            check_run.font.size = Pt(11)
            add_formatted_runs(p, m.group(1), base_size=11)
            i += 1
            continue

        if stripped.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.right_indent = Cm(0.6)
            text = stripped.lstrip('>').strip()
            run = p.add_run(f'"{text}"')
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        para_lines = [stripped]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not re.match(r'^(#{1,6}\s|```|[-*]\s|\d+\.\s|\|)', lines[j].strip()):
            para_lines.append(lines[j].strip())
            j += 1
        text = ' '.join(para_lines)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        add_formatted_runs(p, text, base_size=11)
        i = j

    return doc


def main():
    if not MD_PATH.exists():
        raise FileNotFoundError(f"Markdown não encontrado: {MD_PATH}")

    md_text = MD_PATH.read_text(encoding='utf-8')

    # Cópia temporária do template original
    tmp_template = ROOT / "scripts" / "_template_working.docx"
    shutil.copy(str(SOURCE_DOCX), str(tmp_template))

    # Abre o template
    doc = Document(str(tmp_template))

    # Limpa o corpo (mantém sectPr)
    body = doc.element.body
    for child in list(body):
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)

    # Corrige margens do sectPr (template original tem valores float que quebram)
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # Processa markdown
    process_markdown(md_text, doc)

    # Salva
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    tmp_output = ROOT / "scripts" / "_output.docx"
    doc.save(str(tmp_output))

    # Substitui o arquivo na pasta DOCUMENTAÇÃO
    shutil.move(str(tmp_output), str(OUTPUT_PATH))
    tmp_template.unlink(missing_ok=True)

    print(f"[OK] Arquivo gerado: {OUTPUT_PATH}")
    print(f"   Tamanho: {OUTPUT_PATH.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
